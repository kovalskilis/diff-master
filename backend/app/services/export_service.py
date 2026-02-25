from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from typing import Optional, List, Dict, Tuple
import io
import difflib
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
import uuid
import re

from models.document import (
    Snapshot, ArticleVersion, PatchedFragment, Article, EditTarget
)
from services.llm_service import LLMService

# Try importing rich text support (openpyxl >= 3.1)
try:
    from openpyxl.cell.rich_text import CellRichText, TextBlock
    from openpyxl.cell.text import InlineFont
    HAS_RICH_TEXT = True
except ImportError:
    HAS_RICH_TEXT = False


def extract_change_blocks(before_text: str, after_text: str) -> List[Dict]:
    """
    Compare before and after text line-by-line, extract only changed blocks.
    Each block = {'before': str, 'after': str, 'type': 'replace'|'insert'|'delete'}
    """
    before_lines = (before_text or "").splitlines()
    after_lines = (after_text or "").splitlines()
    
    matcher = difflib.SequenceMatcher(None, before_lines, after_lines)
    
    changes = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        
        before_block = '\n'.join(before_lines[i1:i2]).strip()
        after_block = '\n'.join(after_lines[j1:j2]).strip()
        
        # Skip empty changes
        if not before_block and not after_block:
            continue
        
        # Special handling for 'replace' blocks that might contain both replacement and insertion
        # This happens when we replace text (e.g., "агента;" -> "агента.") AND add a new paragraph
        if tag == 'replace' and before_block and after_block:
            # First, try to find a clear pattern: ".\nВ" or ".\n\nВ" (period + newline + capital letter)
            # This is a strong indicator of a new paragraph being added
            para_start_pattern = re.search(r'\.\s*\n+\s*([А-ЯЁ])', after_block)
            if para_start_pattern:
                # Found pattern indicating new paragraph
                para_start_pos = para_start_pattern.start() + 1  # Position after the period
                # Skip whitespace after period
                while para_start_pos < len(after_block) and after_block[para_start_pos] in ' \n\r\t':
                    para_start_pos += 1
                
                # Check if the text before para_start_pos is similar to before_block (with minor changes)
                replace_after = after_block[:para_start_pos].strip()
                insert_after = after_block[para_start_pos:].strip()
                
                # Verify that replace_after is similar to before_block (with punctuation change)
                if replace_after and insert_after and len(insert_after) > 30:
                    # Check similarity
                    replace_sm = difflib.SequenceMatcher(None, before_block, replace_after)
                    similarity = replace_sm.ratio()
                    
                    # If similarity is high (>0.7), it's likely a replacement + insertion
                    if similarity > 0.7:
                        print(f"[Export] Splitting replace+insert block (pattern match):")
                        print(f"  Before: {before_block[:100]}...")
                        print(f"  Replace after: {replace_after[:100]}...")
                        print(f"  Insert after: {insert_after[:100]}...")
                        # Add replace block
                        changes.append({
                            'before': before_block,
                            'after': replace_after,
                            'type': 'replace'
                        })
                        # Add insert block
                        changes.append({
                            'before': "",  # Will be set to "Норма отсутствует" later
                            'after': insert_after,
                            'type': 'insert'
                        })
                        continue
            
            # Fallback: Use SequenceMatcher to find if after_block contains substantial new text
            sm = difflib.SequenceMatcher(None, before_block, after_block)
            opcodes = sm.get_opcodes()
            
            # Check if there's an 'insert' operation with substantial new text
            # and a 'replace' operation (which might be just punctuation change)
            split_successful = False
            
            # First, find all replace and insert operations
            replace_ops = []
            insert_ops = []
            for op_tag, i1, i2, j1, j2 in opcodes:
                if op_tag == 'replace':
                    replace_ops.append((i1, i2, j1, j2))
                elif op_tag == 'insert':
                    insert_ops.append((j1, j2))
            
            # If we have both replace and insert, try to split them
            if replace_ops and insert_ops:
                # Find the first insert operation that comes after a replace
                # This typically indicates: replace punctuation + insert new paragraph
                for insert_j1, insert_j2 in insert_ops:
                    insert_text = after_block[insert_j1:insert_j2]
                    # Check if this is substantial new text (more than 30 chars, likely a new paragraph)
                    if len(insert_text.strip()) > 30:
                        # Find the last replace operation before this insert
                        # The replace should end before the insert starts
                        replace_after_end = insert_j1
                        
                        # Check if there's a replace operation that ends near insert_j1
                        # Look for the position where before_block ends (with modifications)
                        # We want to find where the original text (with minor changes) ends
                        
                        # Strategy: find the position in after_block that corresponds to the end of before_block
                        # This is where the replacement ends and insertion begins
                        
                        # Use a more precise method: find where before_block content ends in after_block
                        # by looking for the last character that matches between before and after
                        before_normalized = before_block.rstrip(';.').strip()
                        # Find this normalized text in after_block
                        after_normalized_start = after_block.find(before_normalized)
                        if after_normalized_start != -1:
                            # The replacement ends after the normalized text, possibly with punctuation change
                            # Look for the end of the sentence (period or semicolon) after the normalized text
                            search_start = after_normalized_start + len(before_normalized)
                            # Find the next sentence ending (period or semicolon)
                            sentence_end = -1
                            for i in range(search_start, min(search_start + 10, len(after_block))):
                                if after_block[i] in '.;':
                                    sentence_end = i + 1
                                    break
                            
                            # If we found a sentence end, use it as the split point
                            if sentence_end > 0 and sentence_end < insert_j1:
                                # The replacement ends at sentence_end (including punctuation)
                                replace_after = after_block[:sentence_end].strip()
                                # Insert starts after the punctuation
                                # Skip any whitespace/newlines after the punctuation
                                insert_start = sentence_end
                                while insert_start < len(after_block) and after_block[insert_start] in ' \n\r\t':
                                    insert_start += 1
                                insert_after = after_block[insert_start:].strip()
                            else:
                                # Fallback: try to find where the new paragraph starts
                                # Look for a pattern like ".\nВ" or ".\n\nВ" (period followed by newline and capital letter)
                                # This indicates the start of a new paragraph
                                para_start_pattern = re.search(r'\.\s*\n+\s*([А-ЯЁ])', after_block[search_start:])
                                if para_start_pattern:
                                    para_start_pos = search_start + para_start_pattern.start() + 1  # After the period
                                    # Skip whitespace after period
                                    while para_start_pos < len(after_block) and after_block[para_start_pos] in ' \n\r\t':
                                        para_start_pos += 1
                                    replace_after = after_block[:para_start_pos].strip()
                                    insert_after = after_block[para_start_pos:].strip()
                                else:
                                    # Last fallback: use insert_j1 as the split point
                                    replace_after = after_block[:insert_j1].strip()
                                    insert_after = after_block[insert_j1:].strip()
                        else:
                            # Fallback: try to find pattern ".\nВ" or similar
                            para_start_pattern = re.search(r'\.\s*\n+\s*([А-ЯЁ])', after_block)
                            if para_start_pattern:
                                para_start_pos = para_start_pattern.start() + 1
                                while para_start_pos < len(after_block) and after_block[para_start_pos] in ' \n\r\t':
                                    para_start_pos += 1
                                replace_after = after_block[:para_start_pos].strip()
                                insert_after = after_block[para_start_pos:].strip()
                            else:
                                # Last fallback: use insert_j1 as the split point
                                replace_after = after_block[:insert_j1].strip()
                                insert_after = after_block[insert_j1:].strip()
                        
                        # Only split if we have both parts and replace_after is similar to before_block
                        if replace_after and insert_after:
                            # Check if replace_after is similar to before_block (with minor changes)
                            replace_sm = difflib.SequenceMatcher(None, before_block, replace_after)
                            similarity = replace_sm.ratio()
                            
                            # If similarity is high (>0.7), it's likely a replacement + insertion
                            if similarity > 0.7 or (len(replace_after) > 0 and len(insert_after) > 30):
                                print(f"[Export] Splitting replace+insert block:")
                                print(f"  Before: {before_block[:100]}...")
                                print(f"  Replace after: {replace_after[:100]}...")
                                print(f"  Insert after: {insert_after[:100]}...")
                                # Add replace block
                                changes.append({
                                    'before': before_block,
                                    'after': replace_after,
                                    'type': 'replace'
                                })
                                # Add insert block
                                changes.append({
                                    'before': "",  # Will be set to "Норма отсутствует" later
                                    'after': insert_after,
                                    'type': 'insert'
                                })
                                split_successful = True
                                break
                            else:
                                print(f"[Export] Split condition not met: similarity={similarity:.2f}, replace_len={len(replace_after)}, insert_len={len(insert_after)}")
            
            # If we successfully split, skip adding the original block
            if split_successful:
                continue
        
        # Default: add block as-is
        changes.append({
            'before': before_block,
            'after': after_block,
            'type': tag  # 'replace', 'insert', 'delete'
        })
    
    return changes


def build_rich_after_text(before_block: str, after_block: str, change_type: str):
    """
    Build rich text for the 'after' column with bold on changed/added words.
    Returns CellRichText if available, otherwise plain string.
    """
    if not HAS_RICH_TEXT:
        return after_block
    
    try:
        bold_font = InlineFont(b=True)
        
        # Pure insertion — all text is bold
        if change_type == 'insert' or not before_block.strip():
            return CellRichText(TextBlock(bold_font, after_block))
        
        # Replace — find changed words and bold them
        before_words = before_block.split()
        after_words = after_block.split()
        
        matcher = difflib.SequenceMatcher(None, before_words, after_words)
        
        parts = []
        first = True
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            chunk = ' '.join(after_words[j1:j2])
            if not chunk:
                continue
            
            # Add space between chunks (except first)
            prefix = '' if first else ' '
            first = False
            
            if tag == 'equal':
                parts.append(prefix + chunk)
            else:
                parts.append(TextBlock(bold_font, prefix + chunk))
        
        if parts:
            return CellRichText(*parts)
        return after_block
    except Exception as e:
        print(f"[RichText] Fallback to plain text: {e}")
        return after_block


def extract_effective_date_heuristic(before_text: str, after_text: str, instruction: Optional[str]) -> Optional[str]:
    """
    Extract effective date using heuristics (no LLM).
    Searches for dates near 'вступ' keyword.
    """
    context_sources = [after_text or "", instruction or "", before_text or ""]
    
    # 1) dd.mm.yyyy near 'вступ'
    for src in context_sources:
        for m in re.finditer(r'вступ\w{0,12}[^\.]{0,140}?(\d{1,2}\.\d{1,2}\.\d{4})', src, flags=re.IGNORECASE | re.DOTALL):
            return m.group(1)
    
    # 2) Textual russian date: "с 1 января 2026 года"
    month_map = {
        "января": "01", "февраля": "02", "марта": "03", "апреля": "04",
        "мая": "05", "июня": "06", "июля": "07", "августа": "08",
        "сентября": "09", "октября": "10", "ноября": "11", "декабря": "12",
    }
    textual_pattern = re.compile(
        r'вступ\w{0,12}[^\.]{0,140}?с\s+(\d{1,2})\s+(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+(\d{4})',
        flags=re.IGNORECASE
    )
    for src in context_sources:
        t = textual_pattern.search(src)
        if t:
            day = int(t.group(1))
            month = month_map[t.group(2).lower()]
            year = t.group(3)
            return f"{day:02d}.{month}.{year}"
    
    # 3) General date pattern: "по истечении одного месяца со дня официального опубликования"
    for src in context_sources:
        m = re.search(r'по истечении[^\.]{0,100}опубликования', src, flags=re.IGNORECASE)
        if m:
            return "вступает в силу по истечении одного месяца со дня официального опубликования настоящего Федерального закона"
    
    return None


def detect_banking_heuristic(*texts: Optional[str]) -> bool:
    """Simple keyword-based detector for banking-related changes."""
    combined = " ".join([t or "" for t in texts]).lower()
    if not combined.strip():
        return False
    keywords = [
        "банк россии", "центральный банк", "цб рф", "цбр",
        "банк", "банковск", "кредитн", "расчетный счет",
        "корреспондентский счет", "вклад", "депозит",
        "банковская гаранти", "кредитная организация",
        "платежная система", "платежный агент", "ипотек"
    ]
    for kw in keywords:
        if kw in combined:
            return True
    return False


def normalize_paragraph_number(abzac_word: str) -> str:
    """
    Convert ordinal number from prepositional/genitive case to nominative case.
    Examples: "девятом" -> "девятый", "десятом" -> "десятый", "втором" -> "второй"
    """
    # Mapping of common ordinal numbers from various cases to nominative
    case_mapping = {
        # Prepositional case (в каком?) -> Nominative (какой?)
        "первом": "первый",
        "втором": "второй",
        "третьем": "третий",
        "четвертом": "четвертый",
        "пятом": "пятый",
        "шестом": "шестой",
        "седьмом": "седьмой",
        "восьмом": "восьмой",
        "девятом": "девятый",
        "десятом": "десятый",
        "одиннадцатом": "одиннадцатый",
        "двенадцатом": "двенадцатый",
        # Genitive case (какого?) -> Nominative
        "первого": "первый",
        "второго": "второй",
        "третьего": "третий",
        "четвертого": "четвертый",
        "пятого": "пятый",
        "шестого": "шестой",
        "седьмого": "седьмой",
        "восьмого": "восьмой",
        "девятого": "девятый",
        "десятого": "десятый",
        "одиннадцатого": "одиннадцатый",
        "двенадцатого": "двенадцатый",
        # Dative case (какому?) -> Nominative
        "первому": "первый",
        "второму": "второй",
        "третьему": "третий",
        "четвертому": "четвертый",
        "пятому": "пятый",
        "шестому": "шестой",
        "седьмому": "седьмой",
        "восьмому": "восьмой",
        "девятому": "девятый",
        "десятому": "десятый",
        # Instrumental case (каким?) -> Nominative
        "первым": "первый",
        "вторым": "второй",
        "третьим": "третий",
        "четвертым": "четвертый",
        "пятым": "пятый",
        "шестым": "шестой",
        "седьмым": "седьмой",
        "восьмым": "восьмой",
        "девятым": "девятый",
        "десятым": "десятый",
    }
    
    abzac_lower = abzac_word.lower()
    if abzac_lower in case_mapping:
        return case_mapping[abzac_lower]
    
    # If not found in mapping, try to convert common endings
    # -ом, -ем -> -ый, -ий, -ой
    if abzac_lower.endswith('ом') or abzac_lower.endswith('ем'):
        # Remove ending and add nominative ending
        stem = abzac_lower[:-2]
        if stem.endswith('трет'):
            return "третий"
        elif stem.endswith('четверт'):
            return "четвертый"
        elif stem.endswith('пят'):
            return "пятый"
        elif stem.endswith('шест'):
            return "шестой"
        elif stem.endswith('седьм'):
            return "седьмой"
        elif stem.endswith('восьм'):
            return "восьмой"
        elif stem.endswith('девят'):
            return "девятый"
        elif stem.endswith('десят'):
            return "десятый"
        elif stem.endswith('перв'):
            return "первый"
        elif stem.endswith('втор'):
            return "второй"
    
    # If no conversion found, return as-is (might already be in nominative)
    return abzac_word


def build_article_reference(article_number: Optional[str], instruction: Optional[str]) -> str:
    """
    Build a detailed article reference like 'п. 7 ст. 6.1 НК РФ'
    from the instruction text.
    """
    if not article_number:
        return ""
    
    base_ref = f"ст. {article_number} НК РФ"
    
    if instruction:
        # Try to extract specific paragraph/sub-paragraph references
        # e.g. "Пункт 7, подпункт 5:" -> "п. 7 пп. 5 ст. 6.1 НК РФ"
        parts = []
        
        # Look for "Пункт X" or "пункте X"
        punkt_match = re.search(r'[Пп]ункт[еа]?\s+(\d+[\.\d]*)', instruction)
        if punkt_match:
            parts.append(f"п. {punkt_match.group(1)}")
        
        # Look for "подпункт X" or "подпункте X"
        podpunkt_match = re.search(r'[Пп]одпункт[еа]?\s+(\d+[\.\d]*)', instruction)
        if podpunkt_match:
            parts.append(f"пп. {podpunkt_match.group(1)}")
        
        # Look for "абзац X" - normalize to nominative case
        abzac_match = re.search(r'[Аа]бзац[еа]?\s+(\w+)', instruction)
        if abzac_match:
            abzac_num = normalize_paragraph_number(abzac_match.group(1))
            parts.append(f"абз. {abzac_num}")
        
        if parts:
            return f"{' '.join(parts)} {base_ref}"
    
    return base_ref


class ExportService:
    """
    FR-8, FR-9: Export service for text and Excel reports
    """
    
    def __init__(self, session: AsyncSession):
        self.session = session
    
    async def export_as_text(self, snapshot_id: int) -> str:
        """Export snapshot as plain text"""
        result = await self.session.execute(
            select(Snapshot).where(Snapshot.id == snapshot_id)
        )
        snapshot = result.scalar_one_or_none()
        if not snapshot:
            return ""
        
        result = await self.session.execute(
            select(ArticleVersion).where(
                ArticleVersion.snapshot_id == snapshot_id
            )
        )
        versions = result.scalars().all()
        
        text_parts = []
        for version in versions:
            article = version.article
            article_title = f"Статья {article.article_number}" if article else "Статья"
            text_parts.append(f"\n{'='*80}\n")
            text_parts.append(f"{article_title}\n")
            if article and article.title:
                text_parts.append(f"{article.title}\n")
            text_parts.append(f"{'-'*80}\n")
            text_parts.append(f"{version.content}\n")
        
        return ''.join(text_parts)
    
    async def export_as_docx(self, snapshot_id: int) -> bytes:
        """Export snapshot as DOCX"""
        doc = Document()
        doc.add_heading('Экспорт документа', 0)
        
        result = await self.session.execute(
            select(ArticleVersion).where(
                ArticleVersion.snapshot_id == snapshot_id
            )
        )
        versions = result.scalars().all()
        
        for version in versions:
            article = version.article
            article_title = f"Статья {article.article_number}" if article and article.article_number else "Статья"
            doc.add_paragraph(article_title, style='Heading 3')
            if article and article.title:
                doc.add_paragraph(article.title, style='Heading 4')
            doc.add_paragraph(version.content)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    async def export_as_excel(
        self, 
        snapshot_id: Optional[int] = None,
        workspace_file_id: Optional[int] = None,
        user_id: Optional[uuid.UUID] = None
    ) -> bytes:
        """
        FR-9: Export changes as Excel report matching the legal document format.
        
        Columns:
        1. ДЕЙСТВУЮЩАЯ НОРМА НК РФ — only changed paragraphs (before)
        2. НОВАЯ НОРМА — only changed paragraphs (after), with bold on changes
        3. ИЗМЕНЯЕМАЯ/ВВОДИМАЯ НОРМА И ДАТА ВСТУПЛЕНИЯ В ДЕЙСТВИЕ
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "Изменения"
        
        # ── Styles ──
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11, name="Times New Roman")
        data_font = Font(size=10, name="Times New Roman")
        chapter_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
        chapter_font = Font(bold=True, size=10, name="Times New Roman")
        italic_font = Font(italic=True, size=10, name="Times New Roman", color="888888")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # ── Headers (row 1) ──
        headers = [
            "ДЕЙСТВУЮЩАЯ НОРМА НК РФ",
            "НОВАЯ НОРМА",
            "ИЗМЕНЯЕМАЯ/ ВВОДИМАЯ\nНОРМА\nИ\nДАТА ВСТУПЛЕНИЯ В\nДЕЙСТВИЕ",
            "КОММЕНТАРИИ",
            "БАНКОВСКИЙ\nСЕГМЕНТ"
        ]
        
        num_cols = len(headers)
        
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = thin_border
        
        # Column widths
        ws.column_dimensions['A'].width = 55
        ws.column_dimensions['B'].width = 55
        ws.column_dimensions['C'].width = 30
        ws.column_dimensions['D'].width = 45
        ws.column_dimensions['E'].width = 18
        
        # Row height for header
        ws.row_dimensions[1].height = 60
        
        # ── Get patched fragments ──
        query = select(PatchedFragment)
        
        if workspace_file_id:
            target_ids_result = await self.session.execute(
                select(EditTarget.id).where(
                    EditTarget.workspace_file_id == workspace_file_id
                )
            )
            target_ids = [row[0] for row in target_ids_result]
            if target_ids:
                query = query.where(PatchedFragment.edit_target_id.in_(target_ids))
            else:
                buffer = io.BytesIO()
                wb.save(buffer)
                buffer.seek(0)
                return buffer.getvalue()
        
        result = await self.session.execute(query)
        fragments = result.scalars().all()
        
        # ── Sort fragments by article number ──
        async def get_article_for_fragment(frag):
            if frag.article_id:
                r = await self.session.execute(
                    select(Article).where(Article.id == frag.article_id)
                )
                return r.scalar_one_or_none()
            return None
        
        # Preload articles
        fragment_articles = {}
        for frag in fragments:
            article = await get_article_for_fragment(frag)
            fragment_articles[frag.id] = article
        
        # Sort by article number (numeric sort)
        def sort_key(frag):
            art = fragment_articles.get(frag.id)
            if art and art.article_number:
                try:
                    # Handle "6.1" -> (6, 1), "11" -> (11, 0)
                    parts = art.article_number.split('.')
                    return tuple(int(p) for p in parts)
                except ValueError:
                    return (999, 0)
            return (999, 0)
        
        fragments = sorted(fragments, key=sort_key)
        
        # ── Chapter header row ──
        current_row = 2
        
        # Initialize LLM service for comments
        llm_service = LLMService()
        
        # Add "1 Часть НК РФ" header
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=num_cols)
        cell = ws.cell(row=current_row, column=1, value="1 Часть НК РФ")
        cell.fill = chapter_fill
        cell.font = chapter_font
        cell.alignment = Alignment(horizontal="left", vertical="center")
        for col in range(1, num_cols + 1):
            ws.cell(row=current_row, column=col).border = thin_border
        current_row += 1
        
        # ── Fill data rows ──
        for frag_idx, fragment in enumerate(fragments):
            article = fragment_articles.get(fragment.id)
            
            # Get instruction text
            instruction = None
            try:
                if fragment.metadata_json and isinstance(fragment.metadata_json, dict):
                    instruction = fragment.metadata_json.get("instruction")
            except Exception:
                instruction = None
            
            # Extract only changed blocks (not full article)
            change_blocks = extract_change_blocks(
                fragment.before_text or "",
                fragment.after_text or ""
            )
            
            print(f"[Export] Fragment {frag_idx}: extracted {len(change_blocks)} change blocks")
            for i, blk in enumerate(change_blocks):
                print(f"  Block {i+1}: type={blk['type']}, before_len={len(blk.get('before', ''))}, after_len={len(blk.get('after', ''))}")
                if blk['type'] == 'insert':
                    print(f"    Insert content preview: {blk.get('after', '')[:100]}...")
            
            if not change_blocks:
                # No changes detected - skip or show as-is
                continue
            
            # Extract effective date (same for all blocks in this fragment)
            effective_date = extract_effective_date_heuristic(
                fragment.before_text or "",
                fragment.after_text or "",
                instruction
            )
            
            # ── Create a separate row for each change block ──
            for block_idx, block in enumerate(change_blocks):
                # Determine which paragraph/sub-article this block belongs to
                # by analyzing the instruction and block content
                block_article_ref = build_article_reference(
                    article.article_number if article else None,
                    instruction
                )
                
                # Try to extract more specific reference for this block
                # by analyzing the instruction text for paragraph numbers
                if instruction and article:
                    # Special handling for insert blocks (new paragraphs)
                    if block['type'] == 'insert':
                        print(f"[Export] Processing insert block {block_idx+1}/{len(change_blocks)}")
                        print(f"  Instruction preview: {instruction[:200] if instruction else 'None'}...")
                        
                        # Strategy 1: Look for "дополнить абзацем [номер] пункта [номер]"
                        # e.g., "дополнить абзацем десятым пункта 17.2"
                        insert_match = re.search(
                            r'(?:дополнить|добавить)\s+абзац[ем]?\s+(\w+)\s+пункт[ае]?\s+(\d+[\.\-\d]*)',
                            instruction,
                            re.IGNORECASE
                        )
                        if insert_match:
                            abzac_num_raw = insert_match.group(1)  # e.g., "десятым" or "десятый"
                            abzac_num = normalize_paragraph_number(abzac_num_raw)  # Convert to nominative
                            punkt_num = insert_match.group(2)  # e.g., "17.2"
                            base_ref = f"ст. {article.article_number} НК РФ"
                            block_article_ref = f"п. {punkt_num} абз. {abzac_num} {base_ref}"
                            print(f"  Found via Strategy 1: {block_article_ref}")
                        else:
                            # Strategy 2: Look for "дополнить абзацем [номер]" (without explicit point)
                            # Then find point number separately
                            # Make sure we find the абзац number that's directly after "дополнить/добавить"
                            # Pattern: "дополнить абзацем десятым" - the number should be within 5 words after "дополнить абзацем"
                            insert_abzac_match = re.search(
                                r'(?:дополнить|добавить)\s+абзац[ем]?\s+(\w+)',
                                instruction,
                                re.IGNORECASE
                            )
                            punkt_matches = re.findall(r'[Пп]ункт[еа]?\s+(\d+[\.\-\d]*)', instruction)
                            
                            print(f"  Strategy 2: insert_abzac_match={insert_abzac_match is not None}, punkt_matches={punkt_matches}")
                            if insert_abzac_match:
                                print(f"  Strategy 2: Found абзац '{insert_abzac_match.group(1)}'")
                            
                            if punkt_matches and insert_abzac_match:
                                # Found both paragraph number and point
                                # For insert blocks, we want the point that's mentioned in context with "дополнить"
                                # Try to find the point number closest to "дополнить абзацем"
                                punkt_num = None
                                insert_pos = insert_abzac_match.start()
                                # Find all point mentions with their positions
                                punkt_positions = []
                                for match in re.finditer(r'[Пп]ункт[еа]?\s+(\d+[\.\-\d]*)', instruction):
                                    punkt_positions.append((match.start(), match.group(1)))
                                
                                # Use the point that appears closest to the insert instruction
                                if punkt_positions:
                                    # Find point closest to insert_pos
                                    closest = min(punkt_positions, key=lambda x: abs(x[0] - insert_pos))
                                    punkt_num = closest[1]
                                else:
                                    punkt_num = punkt_matches[0]  # Fallback to first mentioned point
                                
                                abzac_num_raw = insert_abzac_match.group(1)
                                abzac_num = normalize_paragraph_number(abzac_num_raw)  # Convert to nominative
                                base_ref = f"ст. {article.article_number} НК РФ"
                                block_article_ref = f"п. {punkt_num} абз. {abzac_num} {base_ref}"
                                print(f"  Found via Strategy 2: {block_article_ref}")
                            elif punkt_matches:
                                # Only point found, try to find paragraph number in context
                                punkt_num = punkt_matches[0]
                                base_ref = f"ст. {article.article_number} НК РФ"
                                
                                # Strategy 3: Look for "дополнить абзацем [номер]" or "добавить абзац [номер]"
                                # Pattern 1: "дополнить абзацем десятым" (tvoritelny case)
                                insert_context_match = re.search(
                                    r'(?:дополнить|добавить)\s+абзац[ем]?\s+(\w+)',
                                    instruction,
                                    re.IGNORECASE
                                )
                                if not insert_context_match:
                                    # Pattern 2: "дополнить/добавить" followed by "абзац [номер]" within 30 chars
                                    # This ensures we find the абзац number that's directly related to "дополнить"
                                    insert_context_match = re.search(
                                        r'(?:дополнить|добавить).{0,30}?[Аа]бзац[еа]?\s+(\w+)',
                                        instruction,
                                        re.IGNORECASE | re.DOTALL
                                    )
                                
                                if insert_context_match:
                                    abzac_num_raw = insert_context_match.group(1)
                                    # Verify it's an ordinal number (not a regular word)
                                    if abzac_num_raw.lower() in ['первый', 'первом', 'первого', 'первому', 'первым',
                                                                  'второй', 'втором', 'второго', 'второму', 'вторым',
                                                                  'третий', 'третьем', 'третьего', 'третьему', 'третьим',
                                                                  'четвертый', 'четвертом', 'четвертого', 'четвертому', 'четвертым',
                                                                  'пятый', 'пятом', 'пятого', 'пятому', 'пятым',
                                                                  'шестой', 'шестом', 'шестого', 'шестому', 'шестым',
                                                                  'седьмой', 'седьмом', 'седьмого', 'седьмому', 'седьмым',
                                                                  'восьмой', 'восьмом', 'восьмого', 'восьмому', 'восьмым',
                                                                  'девятый', 'девятом', 'девятого', 'девятому', 'девятым',
                                                                  'десятый', 'десятом', 'десятого', 'десятому', 'десятым',
                                                                  'одиннадцатый', 'одиннадцатом', 'одиннадцатого', 'одиннадцатому', 'одиннадцатым',
                                                                  'двенадцатый', 'двенадцатом', 'двенадцатого', 'двенадцатому', 'двенадцатым']:
                                        abzac_num = normalize_paragraph_number(abzac_num_raw)
                                        block_article_ref = f"п. {punkt_num} абз. {abzac_num} {base_ref}"
                                        print(f"  Found via Strategy 3: {block_article_ref}")
                                    else:
                                        # Not a valid ordinal, try Strategy 4
                                        insert_context_match = None
                                
                                if not insert_context_match:
                                    # Strategy 4: Find all occurrences of "дополнить/добавить" and look for "абзац [номер]" 
                                    # in the pattern "дополнить абзацем [номер]" or "добавить абзац [номер]"
                                    # Search within 50 chars to catch "дополнить абзацем десятым"
                                    insert_positions = []
                                    for match in re.finditer(r'(?:дополнить|добавить)', instruction, re.IGNORECASE):
                                        insert_positions.append(match.end())
                                    
                                    found_abzac = None
                                    for insert_pos in insert_positions:
                                        # Look for pattern "абзац[ем]? [номер]" within 50 chars after "дополнить/добавить"
                                        # This catches "дополнить абзацем десятым"
                                        search_text = instruction[insert_pos:insert_pos+50]
                                        # Pattern 1: "абзацем десятым" or "абзац десятым"
                                        abzac_after = re.search(r'[Аа]бзац[еа]?м?\s+(\w+)', search_text, re.IGNORECASE)
                                        if abzac_after:
                                            abzac_num_raw = abzac_after.group(1)
                                            # Verify it's an ordinal number
                                            if abzac_num_raw.lower() in ['первый', 'первом', 'первого', 'первому', 'первым',
                                                                          'второй', 'втором', 'второго', 'второму', 'вторым',
                                                                          'третий', 'третьем', 'третьего', 'третьему', 'третьим',
                                                                          'четвертый', 'четвертом', 'четвертого', 'четвертому', 'четвертым',
                                                                          'пятый', 'пятом', 'пятого', 'пятому', 'пятым',
                                                                          'шестой', 'шестом', 'шестого', 'шестому', 'шестым',
                                                                          'седьмой', 'седьмом', 'седьмого', 'седьмому', 'седьмым',
                                                                          'восьмой', 'восьмом', 'восьмого', 'восьмому', 'восьмым',
                                                                          'девятый', 'девятом', 'девятого', 'девятому', 'девятым',
                                                                          'десятый', 'десятом', 'десятого', 'десятому', 'десятым',
                                                                          'одиннадцатый', 'одиннадцатом', 'одиннадцатого', 'одиннадцатому', 'одиннадцатым',
                                                                          'двенадцатый', 'двенадцатом', 'двенадцатого', 'двенадцатому', 'двенадцатым']:
                                                found_abzac = abzac_num_raw
                                                print(f"  Strategy 4: Found абзац '{abzac_num_raw}' after 'дополнить/добавить'")
                                                break
                                    
                                    if found_abzac:
                                        abzac_num = normalize_paragraph_number(found_abzac)
                                        block_article_ref = f"п. {punkt_num} абз. {abzac_num} {base_ref}"
                                        print(f"  Found via Strategy 4: {block_article_ref}")
                                    else:
                                        # No valid абзац found after "дополнить/добавить"
                                        block_article_ref = f"п. {punkt_num} {base_ref}"
                                        print(f"  No valid абзац found after 'дополнить/добавить', using: {block_article_ref}")
                            else:
                                # No point found, use base reference
                                base_ref = f"ст. {article.article_number} НК РФ"
                                if insert_abzac_match:
                                    abzac_num_raw = insert_abzac_match.group(1)
                                    abzac_num = normalize_paragraph_number(abzac_num_raw)
                                    block_article_ref = f"абз. {abzac_num} {base_ref}"
                                    print(f"  Found via Strategy 6: {block_article_ref}")
                                else:
                                    block_article_ref = base_ref
                                    print(f"  No matches, using base: {block_article_ref}")
                    else:
                        # For replace/delete blocks, use standard logic
                        punkt_matches = re.findall(r'[Пп]ункт[еа]?\s+(\d+[\.\-\d]*)', instruction)
                        if punkt_matches:
                            # Try to match blocks to paragraphs based on content
                            punkt_num = None
                            if len(punkt_matches) == len(change_blocks):
                                # Perfect match: one paragraph per block
                                punkt_num = punkt_matches[block_idx]
                            elif len(punkt_matches) > 0:
                                # Multiple paragraphs, but fewer or more blocks
                                # Try to match by checking if block content contains paragraph number
                                block_text = (block.get('after', '') or block.get('before', '')).lower()
                                for pnum in punkt_matches:
                                    if pnum.replace('.', r'\.').replace('-', r'\-') in block_text or \
                                       f"пункт {pnum}" in block_text.lower() or \
                                       f"п. {pnum}" in block_text.lower():
                                        punkt_num = pnum
                                        break
                                
                                # If no match found, use index-based matching
                                if not punkt_num:
                                    if block_idx < len(punkt_matches):
                                        punkt_num = punkt_matches[block_idx]
                                    else:
                                        punkt_num = punkt_matches[-1]  # Use last paragraph
                            
                            if punkt_num:
                                # Build specific reference
                                base_ref = f"ст. {article.article_number} НК РФ"
                                
                                # Try to extract абзац number from instruction
                                # Look for patterns like "в абзаце втором", "абзац второй", etc.
                                # First, try to find all абзац mentions in instruction
                                abzac_matches = re.findall(r'[Аа]бзац[еа]?\s+(\w+)', instruction)
                                if abzac_matches:
                                    # If we have multiple blocks and multiple абзац mentions,
                                    # try to match by index
                                    if len(abzac_matches) == len(change_blocks):
                                        abzac_num_raw = abzac_matches[block_idx]
                                        abzac_num = normalize_paragraph_number(abzac_num_raw)  # Convert to nominative
                                        block_article_ref = f"п. {punkt_num} абз. {abzac_num} {base_ref}"
                                    else:
                                        # Use first match for first block, last match for last block
                                        if block_idx == 0:
                                            abzac_num_raw = abzac_matches[0]
                                        elif block_idx >= len(abzac_matches):
                                            abzac_num_raw = abzac_matches[-1]
                                        else:
                                            abzac_num_raw = abzac_matches[block_idx]
                                        abzac_num = normalize_paragraph_number(abzac_num_raw)  # Convert to nominative
                                        block_article_ref = f"п. {punkt_num} абз. {abzac_num} {base_ref}"
                                else:
                                    block_article_ref = f"п. {punkt_num} {base_ref}"
                
                # Build cell content for this block
                if block['type'] == 'delete':
                    before_cell_text = block['before']
                    after_cell_text = ""
                elif block['type'] == 'insert':
                    before_cell_text = "Норма отсутствует"
                    after_cell_text = block['after']
                else:  # replace
                    before_cell_text = block['before']
                    after_cell_text = block['after']
                
                # Combine reference + date for column 3
                col3_parts = []
                if block_article_ref:
                    col3_parts.append(block_article_ref)
                if effective_date:
                    col3_parts.append("")  # empty line separator
                    col3_parts.append(effective_date)
                col3_text = "\n".join(col3_parts)
                
                # ── Write row for this block ──
                # Column 1: ДЕЙСТВУЮЩАЯ НОРМА
                cell_before = ws.cell(row=current_row, column=1)
                if not before_cell_text or before_cell_text == "Норма отсутствует":
                    cell_before.value = "Норма отсутствует"
                    cell_before.font = italic_font
                else:
                    cell_before.value = before_cell_text
                    cell_before.font = data_font
                cell_before.alignment = Alignment(vertical="top", wrap_text=True)
                cell_before.border = thin_border
                
                # Column 2: НОВАЯ НОРМА (with bold on changes)
                cell_after = ws.cell(row=current_row, column=2)
                
                # Build rich text for this single block
                if HAS_RICH_TEXT:
                    rich_value = build_rich_after_text(
                        block['before'], block['after'], block['type']
                    )
                    cell_after.value = rich_value
                else:
                    cell_after.value = after_cell_text
                
                cell_after.font = data_font
                cell_after.alignment = Alignment(vertical="top", wrap_text=True)
                cell_after.border = thin_border
                
                # Column 3: ИЗМЕНЯЕМАЯ НОРМА + ДАТА
                if block['type'] == 'insert':
                    print(f"[Export] INSERT block {block_idx+1}: block_article_ref='{block_article_ref}', col3_text='{col3_text}'")
                cell_ref = ws.cell(row=current_row, column=3, value=col3_text)
                cell_ref.font = data_font
                cell_ref.alignment = Alignment(vertical="top", wrap_text=True)
                cell_ref.border = thin_border
                
                # Column 4: КОММЕНТАРИИ (LLM summary)
                comment_text = ""
                try:
                    comment_text = await llm_service.summarize_edit(
                        before_text=before_cell_text,
                        after_text=after_cell_text,
                        article_number=article.article_number if article else None,
                        instruction=instruction,
                        block_type=block['type']
                    )
                except Exception:
                    if instruction:
                        comment_text = f"Кратко: {instruction[:200]}"
                    else:
                        comment_text = "Изменение текста статьи"
                
                cell_comment = ws.cell(row=current_row, column=4, value=comment_text)
                cell_comment.font = data_font
                cell_comment.alignment = Alignment(vertical="top", wrap_text=True)
                cell_comment.border = thin_border
                
                # Column 5: БАНКОВСКИЙ СЕГМЕНТ (heuristic)
                is_banking = detect_banking_heuristic(
                    instruction, fragment.before_text, fragment.after_text
                )
                cell_banking = ws.cell(row=current_row, column=5, value="Да" if is_banking else "Нет")
                cell_banking.font = data_font
                cell_banking.alignment = Alignment(horizontal="center", vertical="top", wrap_text=True)
                cell_banking.border = thin_border
                
                # Auto row height (estimate based on content length)
                max_len = max(len(before_cell_text), len(after_cell_text), 1)
                estimated_lines = max_len // 60 + (before_cell_text.count('\n') if before_cell_text else 0) + 2
                ws.row_dimensions[current_row].height = max(30, min(estimated_lines * 15, 400))
                
                print(f"[Export] Row {current_row}: art={article.article_number if article else '—'}, "
                      f"block={block_idx+1}/{len(change_blocks)}, type={block['type']}, "
                      f"before_len={len(before_cell_text)}, after_len={len(after_cell_text)}")
                
                current_row += 1
        
        # ── Freeze header row ──
        ws.freeze_panes = "A2"
        
        # ── Print area ──
        ws.print_area = f"A1:E{current_row - 1}"
        ws.page_setup.orientation = 'landscape'
        ws.page_setup.fitToWidth = 1
        ws.page_setup.fitToHeight = 0
        
        # Save
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
