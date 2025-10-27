import io
import re
from typing import Dict, Optional, List
from docx import Document


def parse_document_structure(docx_content: bytes) -> Optional[Dict[str, Dict[str, str]]]:
    """
    Parse DOCX document structure to extract articles with their content.
    
    Args:
        docx_content: Raw bytes of the DOCX file
        
    Returns:
        Dictionary where keys are article numbers (e.g., "1", "11.3") and values are:
        {
            "title": "Article title",
            "content": "Full article text including all paragraphs"
        }
        Returns None if parsing fails.
    """
    try:
        doc = Document(io.BytesIO(docx_content))
        articles = {}
        
        # Pattern to match article headers: "Статья 1. Title" or "Статья 11.3. Title"
        article_pattern = re.compile(r'^Статья\s+(\d+(?:\.\d+)?)\.\s*(.*)', re.IGNORECASE)
        
        current_article = None
        current_content = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # Skip service lines (ConsultantPlus, etc.)
            if any(service in text.lower() for service in ['консультантплюс', 'consultantplus', '©', 'copyright']):
                continue
                
            # Check if this is an article header
            match = article_pattern.match(text)
            if match:
                # Save previous article if exists
                if current_article:
                    articles[current_article['number']] = {
                        'title': current_article['title'],
                        'content': '\n'.join(current_content).strip()
                    }
                
                # Start new article
                article_number = match.group(1)
                article_title = match.group(2)
                current_article = {
                    'number': article_number,
                    'title': article_title
                }
                current_content = [text]  # Include the header line
            else:
                # Add content to current article
                if current_article:
                    current_content.append(text)
        
        # Save the last article
        if current_article:
            articles[current_article['number']] = {
                'title': current_article['title'],
                'content': '\n'.join(current_content).strip()
            }
        
        return articles if articles else None
        
    except Exception as e:
        print(f"Error parsing document structure: {e}")
        return None


def parse_txt_structure(txt_content: str) -> Optional[Dict[str, Dict[str, str]]]:
    """
    Parse TXT document structure to extract articles with their content.
    Similar to DOCX parsing but works with plain text.
    
    Args:
        txt_content: Plain text content of the document
        
    Returns:
        Dictionary where keys are article numbers and values contain title and content.
        Returns None if parsing fails.
    """
    try:
        articles = {}
        
        # Pattern to match article headers
        article_pattern = re.compile(r'^Статья\s+(\d+(?:\.\d+)?)\.\s*(.*)', re.IGNORECASE)
        
        current_article = None
        current_content = []
        
        for line in txt_content.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            # Skip service lines
            if any(service in line.lower() for service in ['консультантплюс', 'consultantplus', '©', 'copyright']):
                continue
                
            # Check if this is an article header
            match = article_pattern.match(line)
            if match:
                # Save previous article if exists
                if current_article:
                    articles[current_article['number']] = {
                        'title': current_article['title'],
                        'content': '\n'.join(current_content).strip()
                    }
                
                # Start new article
                article_number = match.group(1)
                article_title = match.group(2)
                current_article = {
                    'number': article_number,
                    'title': article_title
                }
                current_content = [line]  # Include the header line
            else:
                # Add content to current article
                if current_article:
                    current_content.append(line)
        
        # Save the last article
        if current_article:
            articles[current_article['number']] = {
                'title': current_article['title'],
                'content': '\n'.join(current_content).strip()
            }
        
        return articles if articles else None
        
    except Exception as e:
        print(f"Error parsing TXT document structure: {e}")
        return None


def extract_edits_for_review(content, file_type: str = "file") -> Dict[str, str]:
    """
    Extract edits by article for user review and approval
    
    Args:
        content: Raw content of the edits file (bytes or str)
        file_type: Type of file ("docx", "txt", or "file")
        
    Returns:
        Dictionary with article numbers as keys and full article content as values
        {
            "1": "Статья 1\nВнести в часть первую Налогового кодекса...\n1) в пункте 7 статьи 6.1:\nа) слова 'рабочий день'...",
            "2": "Статья 2\n[содержимое статьи 2]",
            "3": "Статья 3\n[содержимое статьи 3]"
        }
    """
    try:
        print(f"[Parsing] Extracting edits for review, file_type={file_type}, content_type={type(content)}")
        
        # Handle different content types
        if isinstance(content, bytes):
            # Check if this looks like a binary file (e.g., .docx)
            if b'[Content_Types].xml' in content or b'word/' in content or b'PK' in content[:4]:
                print(f"[Parsing] Binary file detected (likely .docx), attempting to parse as DOCX")
                try:
                    text_content = _extract_text_from_docx(content)
                except Exception as e:
                    print(f"[Parsing] Failed to parse as DOCX: {e}")
                    return {"unknown": f"Не удалось обработать файл .docx: {str(e)}"}
            else:
                # Try different encodings for text files
                text_content = None
                for encoding in ['utf-8', 'cp1251', 'windows-1251', 'latin1']:
                    try:
                        text_content = content.decode(encoding)
                        # Check if this looks like a text file with articles
                        if 'Статья' in text_content or 'статья' in text_content:
                            print(f"[Parsing] Successfully decoded with {encoding}")
                            break
                    except UnicodeDecodeError:
                        continue
                
                if text_content is None:
                    # If no encoding worked, use latin1 with error replacement
                    text_content = content.decode('latin1', errors='replace')
                    print(f"[Parsing] Used latin1 with error replacement")
        else:
            # content is already a string
            text_content = str(content)
        
        print(f"[Parsing] Extracted text content: {len(text_content)} characters")
        print(f"[Parsing] First 200 chars: {text_content[:200]}...")
        
        # Check if the content looks like legal text with articles
        if not ('Статья' in text_content or 'статья' in text_content or 'Внести' in text_content):
            print(f"[Parsing] Content doesn't appear to contain legal articles")
            return {"unknown": "Файл не содержит правки в формате статей. Пожалуйста, загрузите текстовый файл с правками."}
        
        # Split by articles
        articles = _split_by_articles(text_content)
        
        print(f"[Parsing] Final result: {len(articles)} articles found")
        return articles
        
    except Exception as e:
        print(f"Error extracting edits for review: {e}")
        import traceback
        traceback.print_exc()
        return {"unknown": f"Ошибка парсинга файла: {str(e)}"}


def parse_and_group_edits(content: bytes, file_type: str = "docx") -> Dict[str, List[str]]:
    """
    Parse and group edits by article number (legacy function for backward compatibility)
    
    Args:
        content: Raw content of the edits file
        file_type: Type of file ("docx" or "txt")
        
    Returns:
        Dictionary with article numbers as keys and lists of edits as values
        {
            "5": ["В Статье 5 пункт 2 изложить...", "Дополнить Статью 5 пунктом 4..."],
            "11.3": ["Исключить абзац третий пункта 1 Статьи 11.3..."],
            "unknown": ["Заменить слово 'федеральный' на 'государственный'..."]
        }
    """
    try:
        # Extract text content based on file type
        if file_type.lower() == "docx":
            text_content = _extract_text_from_docx(content)
        else:  # txt
            if isinstance(content, bytes):
                # Try different encodings
                for encoding in ['utf-8', 'cp1251', 'windows-1251', 'latin1']:
                    try:
                        text_content = content.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    # If all encodings fail, use utf-8 with errors='replace'
                    text_content = content.decode('utf-8', errors='replace')
            else:
                text_content = content
        
        # Split into individual edits
        edits = _split_into_edits(text_content)
        
        # Group edits by article
        grouped_edits = _group_edits_by_article(edits)
        
        return grouped_edits
        
    except Exception as e:
        print(f"Error parsing and grouping edits: {e}")
        return {"unknown": [f"Ошибка парсинга файла: {str(e)}"]}


def _extract_text_from_docx(content: bytes) -> str:
    """Extract text content from DOCX file"""
    try:
        doc = Document(io.BytesIO(content))
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)
        
        return '\n'.join(paragraphs)
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""


def _split_by_articles(text: str) -> Dict[str, str]:
    """
    Split text by articles - each article contains all its edits as one fragment
    Now supports all declensions of the word "статья" and includes content before first article mention
    
    Args:
        text: Full text content of the edits file
        
    Returns:
        Dictionary with article numbers as keys and full article content as values
    """
    articles = {}
    
    # Расширенные паттерны для поиска статей с учетом всех склонений
    article_patterns = [
        # Заголовки статей
        r'^Статья\s+(\d+(?:\.\d+)?)',           # "Статья 1", "Статья 11.3"
        r'^статья\s+(\d+(?:\.\d+)?)',           # "статья 1" (lowercase)
        r'^СТАТЬЯ\s+(\d+(?:\.\d+)?)',           # "СТАТЬЯ 1" (uppercase)
        
        # Ссылки на статьи в разных падежах (включая середину строки)
        r'в\s+статье\s+(\d+(?:\.\d+)?)',        # "в статье 11"
        r'В\s+статье\s+(\d+(?:\.\d+)?)',        # "В статье 11"
        r'статьи\s+(\d+(?:\.\d+)?)',            # "статьи 6.1"
        r'Статьи\s+(\d+(?:\.\d+)?)',            # "Статьи 6.1"
        r'статью\s+(\d+(?:\.\d+)?)',            # "статью 11.3"
        r'Статью\s+(\d+(?:\.\d+)?)',            # "Статью 11.3"
        r'статьей\s+(\d+(?:\.\d+)?)',           # "статьей 11.3"
        r'Статьей\s+(\d+(?:\.\d+)?)',           # "Статьей 11.3"
        r'статьях\s+(\d+(?:\.\d+)?)',           # "статьях 11.3"
        r'Статьях\s+(\d+(?:\.\d+)?)',           # "Статьях 11.3"
        r'статей\s+(\d+(?:\.\d+)?)',            # "статей 11.3"
        r'Статей\s+(\d+(?:\.\d+)?)',            # "Статей 11.3"
        
        # Сокращения (включая середину строки)
        r'ст\.\s*(\d+(?:\.\d+)?)',              # "ст. 1"
        r'Ст\.\s*(\d+(?:\.\d+)?)',              # "Ст. 1"
        r'СТ\.\s*(\d+(?:\.\d+)?)',              # "СТ. 1"
        
        # С двоеточием или тире
        r'^Статья\s+(\d+(?:\.\d+)?)\s*[:\-]',    # "Статья 1:", "Статья 1-"
        r'^статья\s+(\d+(?:\.\d+)?)\s*[:\-]',    # "статья 1:", "статья 1-" (lowercase)
        r'^(\d+(?:\.\d+)?)\s*[:\-]',            # "1:", "11.3-" (just number)
        
        # Нумерованные пункты с упоминанием статей
        r'^\d+\)\s+в\s+статье\s+(\d+(?:\.\d+)?)',  # "1) в статье 6.1"
        r'^\d+\)\s+в\s+пункте\s+\d+\s+статьи\s+(\d+(?:\.\d+)?)',  # "1) в пункте 7 статьи 6.1"
        r'^\d+\)\s+в\s+пункте\s+\d+\s+статьи\s+(\d+(?:\.\d+)?)',  # "1) в пункте 2 статьи 11.2"
    ]
    
    # Найти все упоминания статей в тексте
    found_articles = set()
    article_positions = {}  # Позиции начала каждой статьи
    
    for pattern in article_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            article_num = match.group(1)
            found_articles.add(article_num)
            # Сохраняем позицию первого упоминания статьи
            if article_num not in article_positions:
                article_positions[article_num] = match.start()
    
    print(f"[Parsing] Found article references: {sorted(found_articles)}")
    
    if found_articles:
        # Сортируем статьи по позиции в тексте
        sorted_articles = sorted(found_articles, key=lambda x: article_positions[x])
        
        # Разделяем текст по статьям
        for i, article_num in enumerate(sorted_articles):
            start_pos = article_positions[article_num]
            
            # Определяем конец статьи (начало следующей или конец текста)
            if i + 1 < len(sorted_articles):
                next_article = sorted_articles[i + 1]
                end_pos = article_positions[next_article]
            else:
                end_pos = len(text)
            
            # Извлекаем содержимое статьи
            article_content = text[start_pos:end_pos].strip()
            
            # Если это первая статья, включаем весь текст до неё
            if i == 0:
                # Находим начало первого упоминания статьи
                first_mention_pos = article_positions[article_num]
                # Включаем весь текст от начала до конца статьи
                article_content = text[:end_pos].strip()
            
            articles[article_num] = article_content
            print(f"[Parsing] Article {article_num}: {len(article_content)} characters")
    
    # Если статьи не найдены, попробуем найти ссылки на статьи в тексте
    if not articles:
        print(f"[Parsing] No article headers found, searching for article references...")
        articles = _extract_articles_from_text(text)
    
    # Если всё ещё не найдены статьи, обрабатываем весь текст как неизвестный
    if not articles:
        articles["unknown"] = text.strip()
        print(f"[Parsing] No articles found, treating as unknown")
    
    print(f"[Parsing] Found {len(articles)} articles: {list(articles.keys())}")
    return articles


def _extract_articles_from_text(text: str) -> Dict[str, str]:
    """
    Extract articles from text by finding article references when no headers are found
    Now supports all declensions of the word "статья"
    
    Args:
        text: Full text content
        
    Returns:
        Dictionary with article numbers as keys and content as values
    """
    articles = {}
    
    # Расширенные паттерны для поиска ссылок на статьи в тексте
    article_patterns = [
        r'статья\s+(\d+(?:\.\d+)?)',           # "статья 1", "статья 11.3"
        r'Статья\s+(\d+(?:\.\d+)?)',           # "Статья 1", "Статья 11.3"
        r'СТАТЬЯ\s+(\d+(?:\.\d+)?)',           # "СТАТЬЯ 1", "СТАТЬЯ 11.3"
        r'статье\s+(\d+(?:\.\d+)?)',           # "в статье 1"
        r'Статье\s+(\d+(?:\.\d+)?)',           # "В статье 1"
        r'статьи\s+(\d+(?:\.\d+)?)',           # "статьи 1"
        r'Статьи\s+(\d+(?:\.\d+)?)',           # "Статьи 1"
        r'статью\s+(\d+(?:\.\d+)?)',           # "статью 1"
        r'Статью\s+(\d+(?:\.\d+)?)',           # "Статью 1"
        r'статьей\s+(\d+(?:\.\d+)?)',          # "статьей 1"
        r'Статьей\s+(\d+(?:\.\d+)?)',          # "Статьей 1"
        r'статьях\s+(\d+(?:\.\d+)?)',          # "статьях 1"
        r'Статьях\s+(\d+(?:\.\d+)?)',          # "Статьях 1"
        r'статей\s+(\d+(?:\.\d+)?)',           # "статей 1"
        r'Статей\s+(\d+(?:\.\d+)?)',           # "Статей 1"
        r'ст\.\s*(\d+(?:\.\d+)?)',             # "ст. 1"
        r'Ст\.\s*(\d+(?:\.\d+)?)',             # "Ст. 1"
        r'СТ\.\s*(\d+(?:\.\d+)?)',             # "СТ. 1"
    ]
    
    # Найти все номера статей, упомянутых в тексте
    found_articles = set()
    article_positions = {}  # Позиции упоминаний статей
    
    for pattern in article_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            article_num = match.group(1)
            found_articles.add(article_num)
            # Сохраняем позицию первого упоминания
            if article_num not in article_positions:
                article_positions[article_num] = match.start()
    
    # Если найдены ссылки на статьи, создаём записи для них
    if found_articles:
        print(f"[Parsing] Found article references: {sorted(found_articles)}")
        
        # Сортируем статьи по позиции в тексте
        sorted_articles = sorted(found_articles, key=lambda x: article_positions[x])
        
        # Разделяем текст по статьям
        for i, article_num in enumerate(sorted_articles):
            start_pos = article_positions[article_num]
            
            # Определяем конец статьи (начало следующей или конец текста)
            if i + 1 < len(sorted_articles):
                next_article = sorted_articles[i + 1]
                end_pos = article_positions[next_article]
            else:
                end_pos = len(text)
            
            # Извлекаем содержимое статьи
            article_content = text[start_pos:end_pos].strip()
            
            # Если это первая статья, включаем весь текст до неё
            if i == 0:
                article_content = text[:end_pos].strip()
            
            articles[article_num] = article_content
            print(f"[Parsing] Created article {article_num} with {len(article_content)} characters")
    
    return articles


def _split_into_edits(text: str) -> List[str]:
    """
    Split text into individual edits
    
    This function tries to identify individual edits based on common patterns:
    - Each paragraph is a separate edit
    - Items in numbered/bulleted lists are separate edits
    - Lines starting with common edit keywords are separate edits
    """
    edits = []
    
    # Common patterns that indicate start of a new edit
    edit_start_patterns = [
        r'^\d+[\.\)]\s*',  # Numbered lists: "1. ", "2) "
        r'^[а-я]\)\s*',    # Lettered lists: "а) ", "б) "
        r'^[-•]\s*',       # Bullet points: "- ", "• "
        r'^(В\s+статье|Исключить|Дополнить|Заменить|Изложить|Внести)',  # Common edit verbs
    ]
    
    lines = text.split('\n')
    current_edit = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if this line starts a new edit
        is_new_edit = any(re.match(pattern, line, re.IGNORECASE) for pattern in edit_start_patterns)
        
        if is_new_edit and current_edit:
            # Save previous edit
            edit_text = '\n'.join(current_edit).strip()
            if edit_text:
                edits.append(edit_text)
            current_edit = [line]
        else:
            # Continue current edit
            current_edit.append(line)
    
    # Add last edit
    if current_edit:
        edit_text = '\n'.join(current_edit).strip()
        if edit_text:
            edits.append(edit_text)
    
    # If no edits were found using patterns, split by paragraphs
    if not edits:
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        edits = paragraphs
    
    return edits


def _group_edits_by_article(edits: List[str]) -> Dict[str, List[str]]:
    """
    Group edits by article number using regex patterns
    Now supports all declensions of the word "статья"
    """
    grouped_edits = {}
    
    # Расширенные паттерны для поиска ссылок на статьи (упорядочены по специфичности)
    article_patterns = [
        r'статья\s+(\d+(?:\.\d+)?)',           # "статья 5", "статья 11.3"
        r'Статья\s+(\d+(?:\.\d+)?)',           # "Статья 5", "Статья 11.3"
        r'СТАТЬЯ\s+(\d+(?:\.\d+)?)',           # "СТАТЬЯ 5", "СТАТЬЯ 11.3"
        r'статье\s+(\d+(?:\.\d+)?)',           # "в статье 5"
        r'Статье\s+(\d+(?:\.\d+)?)',           # "В статье 5"
        r'статьи\s+(\d+(?:\.\d+)?)',           # "статьи 5"
        r'Статьи\s+(\d+(?:\.\d+)?)',           # "Статьи 5"
        r'статью\s+(\d+(?:\.\d+)?)',           # "статью 5"
        r'Статью\s+(\d+(?:\.\d+)?)',           # "Статью 5"
        r'статьей\s+(\d+(?:\.\d+)?)',          # "статьей 5"
        r'Статьей\s+(\d+(?:\.\d+)?)',          # "Статьей 5"
        r'статьях\s+(\d+(?:\.\d+)?)',          # "статьях 5"
        r'Статьях\s+(\d+(?:\.\d+)?)',          # "Статьях 5"
        r'статей\s+(\d+(?:\.\d+)?)',           # "статей 5"
        r'Статей\s+(\d+(?:\.\d+)?)',           # "Статей 5"
        r'ст\.\s*(\d+(?:\.\d+)?)',             # "ст. 5"
        r'Ст\.\s*(\d+(?:\.\d+)?)',             # "Ст. 5"
        r'СТ\.\s*(\d+(?:\.\d+)?)',             # "СТ. 5"
    ]
    
    for edit in edits:
        article_found = False
        found_articles = set()
        
        # Try to find article numbers in the edit text
        for pattern in article_patterns:
            matches = re.finditer(pattern, edit, re.IGNORECASE)
            for match in matches:
                article_num = match.group(1)
                found_articles.add(article_num)
        
        # Handle multiple articles in one edit
        if len(found_articles) > 1:
            print(f"[Parsing] Edit references multiple articles: {found_articles}")
            # Add to the first found article (could be improved to split the edit)
            article_num = list(found_articles)[0]
            if article_num not in grouped_edits:
                grouped_edits[article_num] = []
            grouped_edits[article_num].append(edit)
            article_found = True
        elif len(found_articles) == 1:
            article_num = list(found_articles)[0]
            if article_num not in grouped_edits:
                grouped_edits[article_num] = []
            grouped_edits[article_num].append(edit)
            article_found = True
        
        # If no article found, add to unknown
        if not article_found:
            if "unknown" not in grouped_edits:
                grouped_edits["unknown"] = []
            grouped_edits["unknown"].append(edit)
    
    # Log statistics
    known_articles = len([k for k in grouped_edits.keys() if k != "unknown"])
    unknown_count = len(grouped_edits.get("unknown", []))
    print(f"[Parsing] Grouped edits: {known_articles} known articles, {unknown_count} unknown")
    
    return grouped_edits
